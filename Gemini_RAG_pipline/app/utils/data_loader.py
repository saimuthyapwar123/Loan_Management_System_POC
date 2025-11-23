import os
import pandas as pd
from PyPDF2 import PdfReader
from docx import Document

def load_text_from_file(filepath: str) -> str:
    """Extracts text content from multiple file types."""
    ext = os.path.splitext(filepath)[1].lower()
    text = ""

    if ext == ".txt":
        with open(filepath, "r", encoding="utf-8") as f:
            text = f.read()

    elif ext == ".pdf":
        reader = PdfReader(filepath)
        text = "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])

    elif ext == ".docx":
        doc = Document(filepath)
        text = "\n".join([p.text for p in doc.paragraphs])

    elif ext == ".pptx":
        doc = Document(filepath)
        text = "\n".join([p.text for p in doc.paragraphs])

    elif ext in [".csv", ".xlsx"]:
        df = pd.read_csv(filepath) if ext == ".csv" else pd.read_excel(filepath)
        text = "\n".join(df.astype(str).apply(lambda x: " | ".join(x), axis=1))

    return text


def load_all_documents(folder_path: str = os.path.join("data", "documents")) -> list:
    """Reads all supported files and returns a list of texts."""
    documents = []
    for filename in os.listdir(folder_path):
        if filename.lower().endswith((".pdf", ".txt", ".docx", ".csv", ".xlsx")):
            filepath = os.path.join(folder_path, filename)
            print(f"📄 Loading: {filename}")
            doc_text = load_text_from_file(filepath)
            if doc_text.strip():
                documents.append(doc_text)
    return documents
